import asyncio
import json
import random
import html
from pathlib import Path
from datetime import datetime, date
import time
import os

# Reply Keyboard va fayl eksporti uchun
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, FSInputFile, ReplyKeyboardRemove
from docx import Document 
from openpyxl import Workbook


import aiohttp
from aiogram import Bot, Dispatcher, Router, F
from aiogram.types import (
    Message, CallbackQuery,
    InlineKeyboardMarkup, InlineKeyboardButton
)
from aiogram.filters import Command
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.exceptions import TelegramAPIError

# ================= CONFIG =================
TOKEN = "8438341822:AAG1rbsBha1KyadxNBJNJ4pZAOiyavj1hLQ" 
DATA_FILE = Path("data.json")
TRANSLATE_API = "https://api.mymemory.translated.net/get"
# ADMIN_PASSWORD fayldan yuklanadi yoki standart qiymat ishlatiladi
ADMIN_PASSWORD_FILE = Path("admin_password.txt")
DEFAULT_ADMIN_PASSWORD = "7777" 
ADMIN_ID = 8438341822 # Admin ID ni o'zingiznikiga o'zgartiring!

# Avtomatik o'yin sozlamalari
GAME_INTERVAL = 12 * 60  # 12 daqiqa sekundda
WORDS_PER_GAME = 10
# =========================================

# Bot va Dispatcher yaratish
bot = Bot(TOKEN)
dp = Dispatcher(storage=MemoryStorage())
router = Router()
dp.include_router(router)

# ================= ADMIN PASSWORD MANAGEMENT =================
def load_admin_password():
    if ADMIN_PASSWORD_FILE.exists():
        return ADMIN_PASSWORD_FILE.read_text("utf-8").strip()
    return DEFAULT_ADMIN_PASSWORD

def save_admin_password(password):
    ADMIN_PASSWORD_FILE.write_text(password, "utf-8")

# Bot ishga tushganda parolni yuklash
ADMIN_PASSWORD = load_admin_password()

# ================= STATES =================
class SearchState(StatesGroup):
    active = State()

class MyDictState(StatesGroup):
    active = State()

class EditState(StatesGroup):
    waiting_word = State() 

class DeleteState(StatesGroup):
    waiting_word = State() 

class GameState(StatesGroup):
    playing = State()
    in_settings = State() 

class AdminState(StatesGroup):
    waiting_password = State()
    viewing_stats = State() 
    waiting_new_password = State() 

class DownloadState(StatesGroup):
    waiting_format = State()
    
class InitialSetupState(StatesGroup): 
    waiting_ui_lang = State()
    waiting_tr_lang = State()

# ================= HELPERS =================
def is_admin(uid):
    return int(uid) == ADMIN_ID # Admin tekshiruvi

# ================= FILE & HISTORY =================
def load_users():
    if not DATA_FILE.exists():
        return {}
    try:
        data = json.loads(DATA_FILE.read_text("utf-8"))
        
        # ***************** YENGI MANTIQ *****************
        for uid in data:
             user_data = data[uid]
             if 'tr_lang' not in user_data:
                 user_data['tr_lang'] = 'uz' 
                 
             # ESKI 'words' ro'yxatini yangi 'dictionaries' strukturasiga o'tkazish (Agar birinchi marta o'tkazilayotgan bo'lsa)
             if 'words' in user_data and user_data['words']:
                 print(f"[{uid}] Eskidan yangi formatga o'tkazilmoqda. (Dictionary separation)")
                 user_data['dictionaries'] = {}
                 for word_item in user_data['words']:
                     # Eski formatda 'tr_code' bo'lmasa, 'uz' deb qabul qilamiz
                     tr_code = word_item.get('tr_code', 'uz') 
                     if tr_code not in user_data['dictionaries']:
                         user_data['dictionaries'][tr_code] = []
                     
                     # Faqat kerakli kalitlarni saqlash (ko, uz)
                     user_data['dictionaries'][tr_code].append({
                         "ko": word_item.get('ko'), 
                         "uz": word_item.get('uz')
                     })
                 del user_data['words'] # Eskisini o'chiramiz
             elif 'dictionaries' not in user_data:
                  # Agar hali umuman lug'at qo'shilmagan bo'lsa, lug'at lug'atini yaratamiz
                  user_data['dictionaries'] = {}
        # ***************** YENGI MANTIQ *****************
                 
        return data
    except json.JSONDecodeError:
        print("XATO: data.json fayli buzilgan. Bo'sh lug'at qaytarildi.")
        return {}


def save_users(d):
    DATA_FILE.write_text(json.dumps(d, ensure_ascii=False, indent=2), "utf-8")

def record_history(uid, action, word):
    """Foydalanuvchi harakatini tarixga yozadi. Endi so'zga tr_code qo'shiladi."""
    users = load_users()
    uid_str = str(uid)
    if uid_str in users:
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        if "history" not in users[uid_str]:
            users[uid_str]["history"] = []
            
        users[uid_str]["history"].append({"action": action, "word": word, "time": now})
        save_users(users)

def get_current_dict(uid):
    """Joriy tanlangan tarjima tiliga mos lug'at ro'yxatini qaytaradi."""
    users = load_users()
    uid_str = str(uid)
    current_tr_lang_code = tr_lang(uid)
    
    if uid_str not in users:
        return []

    # Agar lug'atlar mavjud bo'lmasa, uni yaratish
    if 'dictionaries' not in users[uid_str]:
         users[uid_str]['dictionaries'] = {}
         
    if current_tr_lang_code not in users[uid_str]['dictionaries']:
        users[uid_str]['dictionaries'][current_tr_lang_code] = []
        
    # Lug'atni yuklash va keyinchalik o'zgarish bo'lsa save_users(users) ni chaqirish kerak.
    # Lekin bu funksiya faqat ma'lumotni olish uchun ishlatiladi.
    return users[uid_str]['dictionaries'].get(current_tr_lang_code, [])

def get_all_words_flat(uid):
    """Statistika uchun barcha lug'atlardagi so'zlarni birlashtirib qaytaradi."""
    users = load_users()
    uid_str = str(uid)
    
    if uid_str not in users or 'dictionaries' not in users[uid_str]:
        return []
        
    all_words = []
    
    for tr_code, words_list in users[uid_str]['dictionaries'].items():
        for word_item in words_list:
             # Eksport/Statistika uchun tr_code ni qo'shamiz
             word_item_copy = word_item.copy()
             word_item_copy['tr_code'] = tr_code
             all_words.append(word_item_copy)
             
    return all_words


# ================= EXPORT FUNCTIONS (YANGILANDI) =================
def create_docx(words, uid, tr_lang_code):
    # Faqat joriy lug'at eksport qilinadi (words faqat joriy lug'atdagi so'zlar ro'yxati)
    doc = Document()
    doc.add_heading(f"Lug'at ({tr_lang_code.upper()})", 0)
        
    for w in words:
        translations = ', '.join(w.get('uz', []))
        # w.get('tr_code') endi yo'q, chunki faqat bitta tildan keladi.
        doc.add_paragraph(f"🇰🇷 {w['ko']} -> {tr_lang_code.upper()} {translations}")
    
    file_path = Path(f"dict_{uid}_{tr_lang_code}.docx")
    doc.save(file_path)
    return file_path

def create_xlsx(words, uid, tr_lang_code):
    # Faqat joriy lug'at eksport qilinadi
    wb = Workbook()
    ws = wb.active
    ws.title = f"Lug'at ({tr_lang_code.upper()})"
    ws.append(["Koreyscha So'z", f"Tarjimalar ({tr_lang_code.upper()})"])
    
    for w in words:
        ws.append([w['ko'], ', '.join(w.get('uz', []))])
    
    file_path = Path(f"dict_{uid}_{tr_lang_code}.xlsx")
    wb.save(file_path)
    return file_path

def create_json(words, uid, tr_lang_code):
    # Faqat joriy lug'at eksport qilinadi
    data_to_save = [{"ko": w['ko'], "uz": w.get('uz', [])} for w in words]
    file_path = Path(f"dict_{uid}_{tr_lang_code}.json")
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(data_to_save, f, ensure_ascii=False, indent=4)
    return file_path

# ================= TEXT (tr_lang_code qo'shilgan joylar yangilandi) =================
TEXT = {
    "menu": {
        "uz": "🏠 Asosiy menyu",
        "ru": "🏠 Главное меню",
        "en": "🏠 Main menu",
        "ko": "🏠 메인 메뉴",
        "vi": "🏠 Menu chính"
    },
    "enter_word": {
        "uz": "🇰🇷 So‘z kiriting:",
        "ru": "🇰🇷 Введите слово:",
        "en": "🇰🇷 Enter word:",
        "ko": "🇰🇷 단어 입력:",
        "vi": "🇰🇷 Nhập từ:" 
    },
    "empty": {
        "uz": "📭 Lug‘at bo‘sh. Avval so‘z qo‘shing!",
        "ru": "📭 Пусто. Добавьте слово сначала!",
        "en": "📭 Empty. Add a word first!",
        "ko": "📭 비어 있음. 먼저 단어를 추가하세요!",
        "vi": "📭 Trống. Vui lòng thêm từ trước!" 
    },
    "word_list": {
        "uz": "📘 **Lug‘atingizdagi so‘zlar** ({tr_lang_code}):\n\n", # Kod qo'shildi
        "ru": "📘 **Слова в вашем словаре** ({tr_lang_code}):\n\n",
        "en": "📘 **Words in your dictionary** ({tr_lang_code}):\n\n",
        "ko": "📘 **단어장에 있는 단어** ({tr_lang_code}):\n\n",
        "vi": "📘 **Các từ trong từ điển của bạn** ({tr_lang_code}):\n\n"
    },
    "settings_menu": {
        "uz": "⚙️ Sozlamalar:\n\nTilni o‘zgartirish uchun pastdagi tugmani bosing.",
        "ru": "⚙️ Настройки:\n\nНажмите кнопку ниже, чтобы изменить язык.",
        "en": "⚙️ Settings:\n\nClick the button below to change the language.",
        "ko": "⚙️ 설정:\n\n언어를 변경하려면 아래 버튼을 클릭하세요.",
        "vi": "⚙️ Cài đặt:\n\nNhấp vào nút bên dưới để thay đổi ngôn ngữ." 
    },
    "lang_menu": { 
        "uz": "🌐 Qaysi tilni boshqaruv tili sifatida tanlaysiz?",
        "ru": "🌐 Какой язык вы выберете в качестве языка управления?",
        "en": "🌐 Which language will you choose as the control language?",
        "ko": "🌐 관리 언어로 어떤 언어를 선택하시겠습니까?",
        "vi": "🌐 Bạn sẽ chọn ngôn ngữ nào làm ngôn ngữ điều khiển?"
    },
    "tr_lang_menu": { 
        "uz": "🌐 Koreys tilidan qaysi tilga tarjima qilmoqchisiz?",
        "ru": "🌐 На какой язык вы хотите переводить с корейского?",
        "en": "🌐 Which language do you want to translate from Korean to?",
        "ko": "🌐 한국어에서 어떤 언어로 번역하고 싶으신가요?",
        "vi": "🌐 Bạn muốn dịch từ tiếng Hàn sang ngôn ngữ nào?" 
    },
    "stats_text": {
        "uz": "📊 Sizning statistika:\n\n**Lug‘atdagi jami so‘zlar soni**: {total_words}\n**Joriy lug‘at so‘zlari ({tr_lang_code})**: {current_words}\n\n**Bugungi harakatlar** ({today_date}):\n\nQo‘shilgan: {added}\nTahrirlangan: {edited}\nO‘chirilgan: {deleted}", # O'zgartirildi
        "ru": "📊 Ваша статистика:\n\n**Всего слов в словаре**: {total_words}\n**Слов в текущем словаре ({tr_lang_code})**: {current_words}\n\n**Сегодняшние действия** ({today_date}):\n\nДобавлено: {added}\nИзменено: {edited}\nУдалено: {deleted}",
        "en": "📊 Your Stats:\n\n**Total Words in Dictionary**: {total_words}\n**Words in current dictionary ({tr_lang_code})**: {current_words}\n\n**Today's actions** ({today_date}):\n\nAdded: {added}\nEdited: {edited}\nDeleted: {deleted}",
        "ko": "📊 통계:\n\n**단어장 전체 단어 수**: {total_words}\n**현재 단어장의 단어 ({tr_lang_code})**: {current_words}\n\n**오늘의 활동** ({today_date}):\n\n추가됨: {added}\n수정됨: {edited}\n삭제됨: {deleted}",
        "vi": "📊 Thống kê của bạn:\n\n**Tổng số từ trong Từ điển**: {total_words}\n**Số từ trong từ điển hiện tại ({tr_lang_code})**: {current_words}\n\n**Hoạt động hôm nay** ({today_date}):\n\nĐã thêm: {added}\nĐã chỉnh sửa: {edited}\nĐã xóa: {deleted}" 
    },
    "admin_prompt": {
        "uz": "🔐 Admin paneli.\nParolni kiriting:",
        "ru": "🔐 Панель администратора.\nВведите пароль:",
        "en": "🔐 Admin Panel.\nEnter password:",
        "ko": "🔐 관리자 패널.\n비밀번호를 입력하세요:",
        "vi": "🔐 Bảng điều khiển quản trị.\nNhập mật khẩu:" 
    },
    "admin_menu": {
        "uz": "⭐ Admin paneli. Amaliyotni tanlang:",
        "ru": "⭐ Панель администратора. Выберите действие:",
        "en": "⭐ Admin Panel. Select action:",
        "ko": "⭐ 관리자 패널. 작업을 선택하세요:",
        "vi": "⭐ Bảng điều khiển quản trị. Chọn hành động:" 
    },
    "wrong_password": {
        "uz": "❌ Parol noto‘g‘ri. Menyuga qaytish uchun /start ni bosing.",
        "ru": "❌ Неверный пароль. Нажмите /start для возврата в меню.",
        "en": "❌ Wrong password. Press /start to return to the menu.",
        "ko": "❌ 비밀번호가 틀렸습니다. 메뉴로 돌아가려면 /start를 누르세요.",
        "vi": "❌ Mật khẩu sai. Nhấn /start để quay lại menu." 
    },
    "edit_prompt": {
        "uz": "✏️ Tahrirlash rejimiga kirdingiz. (Faqat joriy lug‘atingizdagilar tahrirlanadi - {tr_lang_code})\n\nTahrirlamoqchi bo‘lgan so‘zni quyidagi formatda kiriting:\n\n`Koreys_so'z.yangi_tarjima1,yangi_tarjima2`\n\n**Masalan:** `사랑.muhabbat,sevgi`\n\nMenyuga qaytish uchun /start ni bosing.", # O'zgartirildi
        "ru": "✏️ Режим редактирования. (Будут редактироваться только слова в вашем текущем словаре - {tr_lang_code})\n\nВведите слово в формате:\n\n`Корейское_слово.новый_перевод1,новый_перевод2`\n\n**Например:** `사랑.любовь,обожание`\n\nДля возврата в меню нажмите /start.",
        "en": "✏️ Edit mode. (Only words in your current dictionary will be edited - {tr_lang_code})\n\nEnter the word in the following format:\n\n`Korean_word.new_translation1,new_translation2`\n\n**Example:** `사랑.love,affection`\n\nTo return to the menu, press /start.",
        "ko": "✏️ 수정 모드입니다. (현재 단어장의 단어만 수정됩니다 - {tr_lang_code})\n\n다음 형식으로 단어를 입력하세요:\n\n`한국어_단어.새_번역1,새_번역2`\n\n**예시:** `사랑.사랑,애정`\n\n메뉴로 돌아가려면 /start를 누르세요.",
        "vi": "✏️ Chế độ chỉnh sửa. (Chỉ những từ trong từ điển hiện tại của bạn sẽ được chỉnh sửa - {tr_lang_code})\n\nNhập từ theo định dạng sau:\n\n`Từ_tiếng_Hàn.bản_dịch_mới1,bản_dịch_mới2`\n\n**Ví dụ:** `사랑.tình yêu,yêu thương`\n\nĐể quay lại menu, nhấn /start." 
    },
    "delete_prompt": {
        "uz": "🗑 O‘chirish rejimiga kirdingiz. (Faqat joriy lug‘atingizdagi so‘zlar o‘chiriladi - {tr_lang_code})\n\nO‘chirmoqchi bo‘lgan koreys so‘zini kiriting.\n\n**Masalan:** `사랑`\n\nMenyuga qaytish uchun /start ni bosing.", # O'zgartirildi
        "ru": "🗑 Режим удаления. (Будут удалены только слова в вашем текущем словаре - {tr_lang_code})\n\nВведите корейское слово, которое хотите удалить.\n\n**Например:** `사랑`\n\nДля возврата в меню нажмите /start.",
        "en": "🗑 Delete mode. (Only words in your current dictionary will be deleted - {tr_lang_code})\n\nEnter the Korean word you want to delete.\n\n**Example:** `사랑`\n\nTo return to the menu, press /start.",
        "ko": "🗑 삭제 모드입니다. (현재 단어장의 단어만 삭제됩니다 - {tr_lang_code})\n\n삭제할 한국어 단어를 입력하세요.\n\n**예시:** `사랑`\n\n메뉴로 돌아가려면 /start를 누르세요.",
        "vi": "🗑 Chế độ xóa. (Chỉ những từ trong từ điển hiện tại của bạn sẽ bị xóa - {tr_lang_code})\n\nNhập từ tiếng Hàn bạn muốn xóa.\n\n**Ví dụ:** `사랑`\n\nĐể quay lại menu, nhấn /start." 
    },
    "edit_success_back_to_mydict": { 
        "uz": "✅ So‘z muvaffaqiyatli tahrirlandi.",
        "ru": "✅ Слово успешно отредактировано.",
        "en": "✅ Word successfully edited.",
        "ko": "✅ 단어가 성공적으로 수정되었습니다.",
        "vi": "✅ Từ đã được chỉnh sửa thành công." 
    },
    "delete_success": {
        "uz": "🗑 So‘z muvaffaqiyatli o‘chirildi. Menyuga qaytish uchun /start ni bosing.",
        "ru": "🗑 Слово успешно удалено. Для возврата в меню нажмите /start.",
        "en": "🗑 Word successfully deleted. To return to the menu, press /start.",
        "ko": "🗑 단어가 성공적으로 삭제되었습니다. 메뉴로 돌아가려면 /start를 누르세요.",
        "vi": "🗑 Từ đã được xóa thành công. Nhấn /start để quay lại menu." 
    },
    "not_found": {
        "uz": "❌ Lug‘atda bunday so‘z topilmadi. Tekshirib qayta kiriting.",
        "ru": "❌ Слово не найдено в словаре. Проверьте и введите снова.",
        "en": "❌ Word not found in the dictionary. Check and re-enter.",
        "ko": "❌ 단어장에서 단어를 찾을 수 없습니다. 확인 후 다시 입력하세요.",
        "vi": "❌ Không tìm thấy từ trong từ điển. Vui lòng kiểm tra và nhập lại." 
    },
    "user_history_text": {
        "uz": "📜 <b>Foydalanuvchi harakatlari tarixi</b> ({username}):\n\n",
        "ru": "📜 <b>История действий пользователя</b> ({username}):\n\n",
        "en": "📜 <b>User Action History</b> ({username}):\n\n",
        "ko": "📜 <b>사용자 활동 기록</b> ({username}):\n\n",
        "vi": "📜 <b>Lịch sử hoạt động của người dùng</b> ({username}):\n\n" 
    },
    "no_history": {
        "uz": "❌ Tarix topilmadi.",
        "ru": "❌ История не найдена.",
        "en": "❌ History not found.",
        "ko": "❌ 기록을 찾을 수 없습니다.",
        "vi": "❌ Không tìm thấy lịch sử." 
    },
    "format_prompt": { 
        "uz": "⬇️ Lug‘at fayli formatini tanlang (Joriy lug‘at: {tr_lang_code}):", # O'zgartirildi
        "ru": "⬇️ Выберите формат файла словаря (Текущий словарь: {tr_lang_code}):",
        "en": "⬇️ Select dictionary file format (Current dictionary: {tr_lang_code}):",
        "ko": "⬇️ 단어장 파일 형식을 선택하세요 (현재 단어장: {tr_lang_code}):",
        "vi": "⬇️ Chọn định dạng tệp từ điển (Từ điển hiện tại: {tr_lang_code}):"
    },
    "file_sent": { 
        "uz": "✅ Fayl yuborildi. Keyingi harakatni tanlang.",
        "ru": "✅ Файл отправлен. Выберите следующее действие.",
        "en": "✅ File sent. Select your next action.",
        "ko": "✅ 파일이 전송되었습니다. 다음 작업을 선택하세요:",
        "vi": "✅ Tệp đã được gửi. Chọn hành động tiếp theo của bạn." 
    },
    "new_password_prompt": { 
        "uz": "🔑 Yangi admin parolini kiriting:",
        "ru": "🔑 Введите новый пароль администратора:",
        "en": "🔑 Enter new admin password:",
        "ko": "🔑 새 관리자 비밀번호를 입력하세요:",
        "vi": "🔑 Nhập mật khẩu quản trị viên mới:" 
    },
    "password_changed": { 
        "uz": "✅ Admin paroli muvaffaqiyatli o‘zgartirildi!",
        "ru": "✅ Пароль администратора успешно изменен!",
        "en": "✅ Admin password successfully changed!",
        "ko": "✅ 관리자 비밀번호가 성공적으로 변경되었습니다!",
        "vi": "✅ Mật khẩu quản trị viên đã được thay đổi thành công!" 
    },
    "game_start": { 
        "uz": "⏳ **Eslatma o‘yini boshlanmoqda!** (Joriy lug‘at: {tr_lang_code}) Har {interval_min} daqiqada {words_count} ta so‘z keladi.\n\nO‘yinni to‘xtatish uchun /stop_game ni bosing. (Birinchi so‘z bir daqiqadan so‘ng keladi.)", # O'zgartirildi
        "ru": "⏳ **Начинается игра-напоминание!** (Текущий словарь: {tr_lang_code}) Каждые {interval_min} минут будет приходить {words_count} слов.\n\nНажмите /stop_game, чтобы остановить игру. (Первое слово придет через минуту.)",
        "en": "⏳ **Reminder game is starting!** (Current dictionary: {tr_lang_code}) {words_count} words will arrive every {interval_min} minutes.\n\nPress /stop_game to stop the game. (First word will arrive in one minute.)",
        "ko": "⏳ **알림 게임이 시작됩니다!** (현재 단어장: {tr_lang_code}) {interval_min}분마다 {words_count}개의 단어가 도착합니다.\n\n게임을 멈추려면 /stop_game을 누르세요. (첫 번째 단어는 1분 후에 도착합니다.)",
        "vi": "⏳ **Trò chơi nhắc nhở sắp bắt đầu!** (Từ điển hiện tại: {tr_lang_code}) {words_count} từ sẽ đến sau mỗi {interval_min} phút.\n\nNhấn /stop_game để dừng trò chơi. (Từ đầu tiên sẽ đến sau một phút.)" 
    },
    "game_stopped": { 
        "uz": "🛑 Eslatma o‘yini to‘xtatildi.",
        "ru": "🛑 Игра-напоминание остановлена.",
        "en": "🛑 Reminder game stopped.",
        "ko": "🛑 알림 게임이 중지되었습니다.",
        "vi": "🛑 Trò chơi nhắc nhở đã dừng." 
    },
    "game_prompt": { 
        "uz": "🎮 **Quyidagi so‘zlarning tarjimasini eslang:** (Tarjima tili: {tr_lang_code})", # O'zgartirildi
        "ru": "🎮 **Вспомните перевод следующих слов:** (Язык перевода: {tr_lang_code})",
        "en": "🎮 **Recall the translation of the following words:** (Translation language: {tr_lang_code})",
        "ko": "🎮 **다음 단어들의 번역을 기억하세요:** (번역 언어: {tr_lang_code})",
        "vi": "🎮 **Nhớ lại bản dịch của các từ sau:** (Ngôn ngữ dịch: {tr_lang_code})"
    },
    "game_reveal": { 
        "uz": "👀 **Tarjimalar:**",
        "ru": "👀 **Переводы:**",
        "en": "👀 **Translations:**",
        "ko": "👀 **번역:**",
        "vi": "👀 **Bản dịch:**" 
    },
    "btn": {
        "search": {"uz": "🔍 Qidirish", "ru": "🔍 Поиск", "en": "🔍 Search", "ko": "🔍 검색", "vi": "🔍 Tìm kiếm"}, 
        "mydict": {"uz": "📘 Lug‘atim", "ru": "📘 Мой словарь", "en": "📘 MyDict", "ko": "📘 단어장", "vi": "📘 Từ điển của tôi"}, 
        "game": {"uz": "🎮 O‘yin", "ru": "🎮 Игра", "en": "🎮 Game", "ko": "🎮 게임", "vi": "🎮 Trò chơi"}, 
        "start_game": {"uz": "▶️ O‘yinni boshlash", "ru": "▶️ Начать игру", "en": "▶️ Start Game", "ko": "▶️ 게임 시작", "vi": "▶️ Bắt đầu trò chơi"}, 
        "stop_game": {"uz": "🛑 O‘yinni to‘xtatish", "ru": "🛑 Остановить игру", "en": "🛑 Stop Game", "ko": "🛑 게임 중지", "vi": "🛑 Dừng trò chơi"}, 
        "stop": {"uz": "🛑 Stop", "ru": "🛑 Стоп", "en": "🛑 Stop", "ko": "🛑 중지", "vi": "🛑 Dừng"}, 
        "edit_all": {"uz": "✏️ Tahrirlash", "ru": "✏️ Изменить", "en": "✏️ Edit", "ko": "✏️ 수정", "vi": "✏️ Chỉnh sửa"}, 
        "delete_all": {"uz": "🗑 O‘chirish", "ru": "🗑 Удалить", "en": "🗑 Delete", "ko": "🗑 삭제", "vi": "🗑 Xóa"}, 
        "reveal": {"uz": "👀 Ko‘rish", "ru": "👀 Показать", "en": "👀 Reveal", "ko": "👀 보기", "vi": "👀 Xem"}, 
        "back_to_menu": {"uz": "◀️ Menyuga", "ru": "◀️ В меню", "en": "◀️ Back to Menu", "ko": "◀️ 메뉴로", "vi": "◀️ Quay lại Menu"}, 
        "cancel": {"uz": "❌ Bekor qilish", "ru": "❌ Отмена", "en": "❌ Cancel", "ko": "❌ 취소", "vi": "❌ Hủy"}, 
        "settings": {"uz": "⚙️ Sozlamalar", "ru": "⚙️ Настройки", "en": "⚙️ Settings", "ko": "⚙️ 설정", "vi": "⚙️ Cài đặt"}, 
        "stats": {"uz": "📊 Statistika", "ru": "📊 Статистика", "en": "📊 Stats", "ko": "📊 통계", "vi": "📊 Thống kê"}, 
        "change_lang": {"uz": "🌐 Boshqaruv tilini o‘zgartirish", "ru": "🌐 Сменить язык", "en": "🌐 Change UI Language", "ko": "🌐 UI 언어 변경", "vi": "🌐 Thay đổi ngôn ngữ UI"}, 
        "change_tr_lang": {"uz": "💬 Tarjima tilini o‘zgartirish", "ru": "💬 Сменить язык перевода", "en": "💬 Change Translation Language", "ko": "💬 번역 언어 변경", "vi": "💬 Thay đổi ngôn ngữ dịch"}, 
        "admin": {"uz": "👑 Admin Panel", "ru": "👑 Админ Панель", "en": "👑 Admin Panel", "ko": "👑 관리자 패널", "vi": "👑 Bảng điều khiển quản trị"}, 
        "change_admin_pass": {"uz": "🔑 Parolni o‘zgartirish", "ru": "🔑 Сменить пароль", "en": "🔑 Change Password", "ko": "🔑 비밀번호 변경", "vi": "🔑 Thay đổi mật khẩu"}, 
        "view_user_history": {"uz": "📜 Tarix", "ru": "📜 История", "en": "📜 History", "ko": "📜 기록", "vi": "📜 Lịch sử"}, 
        "admin_back": {"uz": "◀️ Orqaga", "ru": "◀️ Назад", "en": "◀️ Back", "ko": "◀️ 뒤로", "vi": "◀️ Quay lại"}, 
        "start_btn": {"uz": "➡️ START (Menyu)", "ru": "➡️ СТАРТ (Меню)", "en": "➡️ START (Menu)", "ko": "➡️ 시작 (메뉴)", "vi": "➡️ START (Menu)"}, 
        "download_btn": {"uz": "⬇️ Lug'atni Yuklash", "ru": "⬇️ Скачать Словарь", "en": "⬇️ Download Dictionary", "ko": "⬇️ 단어장 다운로드", "vi": "⬇️ Tải xuống Từ điển"}
    }
}


def lang(uid):
    users = load_users()
    return users.get(str(uid), {}).get("lang", "uz")

def tr_lang(uid): 
    users = load_users()
    return users.get(str(uid), {}).get("tr_lang", "uz")

def t(uid, key, **kwargs):
    # Lug'atdan matnni olish va formatlash (tr_lang_code, interval_min, words_count kabi)
    text = TEXT[key].get(lang(uid), TEXT[key]["uz"])
    return text.format(**kwargs) 

def b(uid, key):
    return TEXT["btn"].get(key, {}).get(lang(uid), TEXT["btn"][key]["uz"]) # Default Uzbek

# ================= KEYBOARDS (O'zgarishsiz qoldi) =================

def lang_kb(prefix="lang"): 
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🇺🇿 Uzbek", callback_data=f"{prefix}_uz")],
        [InlineKeyboardButton(text="🇷🇺 Русский", callback_data=f"{prefix}_ru")],
        [InlineKeyboardButton(text="🇬🇧 English", callback_data=f"{prefix}_en")],
        [InlineKeyboardButton(text="🇻🇳 Vietnamcha", callback_data=f"{prefix}_vi")],
        [InlineKeyboardButton(text="🇰🇷 한국어", callback_data=f"{prefix}_ko")]
    ])

def tr_lang_kb(uid): 
    SUPPORTED_TR_LANGS = [
        ("🇺🇿 Uzbek", "uz"),
        ("🇷🇺 Ruscha", "ru"),
        ("🇬🇧 Inglizcha", "en"),
        ("🇻🇳 Vietnamcha", "vi"),
        ("🇩🇪 Nemischa", "de"),
        ("🇫🇷 Fransuzcha", "fr"),
        ("🇮🇳 Hindcha", "hi"),
        ("🇹🇷 Turkcha", "tr"),
        ("🇹🇭 Tailandcha", "th")
    ]
    
    keyboard = []
    current_tr_lang = tr_lang(uid)
    
    for display_name, code in SUPPORTED_TR_LANGS:
        check = " ✅" if code == current_tr_lang else ""
        keyboard.append([InlineKeyboardButton(text=f"{display_name}{check}", callback_data=f"trlang_{code}")])

    keyboard.append([InlineKeyboardButton(text=b(uid, "admin_back"), callback_data="settings")])
    
    return InlineKeyboardMarkup(inline_keyboard=keyboard)


def main_kb(uid):
    return ReplyKeyboardMarkup(keyboard=[
        [KeyboardButton(text=b(uid, "start_btn"))],
        [KeyboardButton(text=b(uid, "download_btn"))]
    ], resize_keyboard=True, one_time_keyboard=False)

def start_menu_kb_inline(uid):
    keyboard = [
        [InlineKeyboardButton(text=b(uid,"search"), callback_data="search")],
        [InlineKeyboardButton(text=b(uid,"mydict"), callback_data="mydict")],
        [InlineKeyboardButton(text=b(uid,"game"), callback_data="game_menu")],
        [
            InlineKeyboardButton(text=b(uid,"stats"), callback_data="stats"),
            InlineKeyboardButton(text=b(uid,"settings"), callback_data="settings")
        ]
    ]
    if is_admin(uid):
        keyboard.append([InlineKeyboardButton(text=b(uid, "admin"), callback_data="admin_panel")])
        
    return InlineKeyboardMarkup(inline_keyboard=keyboard)

def stop_kb(uid):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=b(uid,"back_to_menu"), callback_data="back_to_inline_menu")] 
    ])

def mydict_management_kb(uid):
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text=b(uid,"edit_all"), callback_data="start_edit"),
            InlineKeyboardButton(text=b(uid,"delete_all"), callback_data="start_delete")
        ],
        [
            InlineKeyboardButton(text=b(uid,"back_to_menu"), callback_data="back_to_inline_menu")
        ]
    ])

def settings_kb(uid):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=b(uid, "change_lang"), callback_data="open_ui_lang_settings")], 
        [InlineKeyboardButton(text=b(uid, "change_tr_lang"), callback_data="open_tr_lang_settings")], 
        [InlineKeyboardButton(text=b(uid,"back_to_menu"), callback_data="back_to_inline_menu")]
    ])

def back_to_menu_kb(uid):
     return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=b(uid,"back_to_menu"), callback_data="back_to_inline_menu")]
    ])

def game_menu_kb(uid): 
     return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=b(uid,"start_game"), callback_data="start_game_auto")],
        [InlineKeyboardButton(text=b(uid,"stop_game"), callback_data="stop_game_auto")],
        [InlineKeyboardButton(text=b(uid,"back_to_menu"), callback_data="back_to_inline_menu")]
    ])
    
def game_reveal_kb(uid):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=b(uid,"reveal"), callback_data="reveal_game_auto")], 
        [InlineKeyboardButton(text=b(uid,"back_to_menu"), callback_data="back_to_inline_menu")]
    ])

def admin_menu_kb(uid): 
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=b(uid, "view_user_history"), callback_data="admin_view_list")],
        [InlineKeyboardButton(text=b(uid, "change_admin_pass"), callback_data="admin_change_pass")],
        [InlineKeyboardButton(text=b(uid, "admin_back"), callback_data="back_to_inline_menu")]
    ])


def admin_user_list_kb(users_data):
    kb = []
    
    sorted_users = sorted(users_data.items(), key=lambda x: int(x[0]))
    
    for uid, data in sorted_users:
        name = data.get("username") or f"ID: {uid}"
        display_name = (name[:30] + '...') if len(name) > 33 else name

        kb.append([InlineKeyboardButton(text=display_name, callback_data=f"view_history:{uid}")])
    
    if users_data:
        sample_uid = list(users_data.keys())[0] 
        kb.append([InlineKeyboardButton(text=b(sample_uid,"admin_back"), callback_data="admin_panel")])
    
    return InlineKeyboardMarkup(inline_keyboard=kb)


def download_kb(uid):
    tr_code = tr_lang(uid).upper()
    keyboard = [[
        InlineKeyboardButton(text=f"PDF (DOCX) - {tr_code}", callback_data="export_docx"),
    ], [
        InlineKeyboardButton(text=f"XLSX (Excel) - {tr_code}", callback_data="export_xlsx"),
        InlineKeyboardButton(text=f"JSON - {tr_code}", callback_data="export_json")
    ], [
        InlineKeyboardButton(text=b(uid, "back_to_menu"), callback_data="back_to_inline_menu")
    ]]
    return InlineKeyboardMarkup(inline_keyboard=keyboard)


# ================= TRANSLATE (O'zgarishsiz qoldi) =================
async def translate(word, tr_code):
    """Koreyscha so'zni foydalanuvchi tanlagan tilga tarjima qilish."""
    lang_pair = f"ko|{tr_code}"
    async with aiohttp.ClientSession() as s:
        async with s.get(TRANSLATE_API, params={"q": word, "langpair": lang_pair}) as r:
            j = await r.json()
            res = []
            
            if "responseData" in j and j["responseData"]["translatedText"]:
                main = html.unescape(j["responseData"]["translatedText"])
                if main:
                    res.append(main)
                
            for m in j.get("matches", []):
                tr = html.unescape(m["translation"])
                if tr and tr not in res:
                    res.append(tr)
            
            if not res:
                res.append("Tarjima topilmadi / No translation found")
            return res

# ================= INITIAL SETUP HANDLERS (O'zgarishsiz qoldi) =================
@router.message(Command("start"))
async def start(m: Message, state: FSMContext):
    uid = str(m.from_user.id)
    await state.clear() 
    users = load_users()

    # Birinchi marta ishga tushirish: UI tilini so'rash
    if uid not in users or "lang" not in users.get(uid, {}):
        
        # Foydalanuvchi ma'lumotlarini saqlash - "words" o'rniga "dictionaries" qo'shildi
        users[uid] = {"lang": None, "tr_lang": "uz", "dictionaries": {}, "history": []} 
        username = m.from_user.username
        if username:
            users[uid]["username"] = f"@{username}"
        elif m.from_user.first_name:
            users[uid]["username"] = m.from_user.first_name
        save_users(users)
        
        await state.set_state(InitialSetupState.waiting_ui_lang)
        await m.answer(TEXT["lang_menu"]["uz"], reply_markup=lang_kb(prefix="setup_uilang"))
    else:
        # Tanish foydalanuvchi: darhol Inline menyuni chiqarish
        await m.answer(t(uid, "menu"), reply_markup=start_menu_kb_inline(uid))

@router.callback_query(F.data.startswith("setup_uilang_"), InitialSetupState.waiting_ui_lang)
async def setup_ui_lang(cb: CallbackQuery, state: FSMContext):
    uid = str(cb.from_user.id)
    users = load_users()
    
    new_lang = cb.data.split("_")[2]
    
    users[uid]["lang"] = new_lang
    save_users(users)
    
    await state.set_state(InitialSetupState.waiting_tr_lang)
    await cb.message.edit_text(t(uid, "tr_lang_menu"), reply_markup=tr_lang_kb(uid))
    await cb.answer(text="✅ Boshqaruv tili o‘rnatildi.")


@router.callback_query(F.data.startswith("trlang_"), InitialSetupState.waiting_tr_lang)
async def setup_tr_lang(cb: CallbackQuery, state: FSMContext):
    uid = str(cb.from_user.id)
    users = load_users()
    
    new_tr_lang = cb.data.split("_")[1]
    
    users[uid]["tr_lang"] = new_tr_lang
    save_users(users)
    
    await state.clear()
    
    await cb.message.edit_text(t(uid,"menu"), reply_markup=None) 
    await cb.message.answer(t(uid,"menu"), reply_markup=start_menu_kb_inline(uid))
    await cb.answer(text="✅ Tarjima tili o‘rnatildi.")

# ================= REPLY KEYBOARD HANDLERS (O'zgarishsiz qoldi) =================

@router.message(F.text.in_([TEXT["btn"]["start_btn"][l] for l in ["uz", "ru", "en", "ko", "vi"]]))
async def handle_start_btn(m: Message, state: FSMContext):
    uid = str(m.from_user.id)
    await state.clear()
    await m.answer(t(uid, "menu"), reply_markup=start_menu_kb_inline(uid))

@router.message(F.text.in_([TEXT["btn"]["download_btn"][l] for l in ["uz", "ru", "en", "ko", "vi"]]))
async def handle_download_btn(m: Message, state: FSMContext):
    uid = str(m.from_user.id)
    
    words_list = get_current_dict(uid) # Faqat joriy lug'atni tekshiramiz

    if not words_list:
        await m.answer(t(uid,"empty"), reply_markup=main_kb(uid))
        return

    await state.set_state(DownloadState.waiting_format)
    tr_code = tr_lang(uid).upper()
    await m.answer(t(uid, "format_prompt", tr_lang_code=tr_code), reply_markup=download_kb(uid))


# ================= INLINE MENU HANDLERS (O'zgarishsiz qoldi) =================

@router.callback_query(F.data=="back_to_inline_menu")
@router.callback_query(F.data=="stop") 
async def back_to_inline_menu(cb:CallbackQuery, state:FSMContext):
    uid = str(cb.from_user.id)
    await state.clear()
    
    try:
        await cb.message.edit_text(t(uid,"menu"), reply_markup=start_menu_kb_inline(uid))
    except TelegramAPIError as e:
        if "message is not modified" not in str(e):
             try:
                 await cb.message.delete()
             except:
                 pass
             await cb.message.answer(t(uid,"menu"), reply_markup=start_menu_kb_inline(uid))

    await cb.answer()

# ================= SEARCH (YANGILANDI) =================
@router.callback_query(F.data=="search")
async def search(cb: CallbackQuery, state:FSMContext):
    uid = str(cb.from_user.id)
    await state.set_state(SearchState.active)
    await cb.message.edit_text(t(uid,"enter_word"), reply_markup=stop_kb(uid))
    await cb.answer()

@router.message(SearchState.active)
async def search_word(m:Message):
    uid = str(m.from_user.id)
    users = load_users()
    word = m.text.strip()
    tr_code = tr_lang(uid) 
    tr = await translate(word, tr_code)
    
    # Joriy lug'atni olish (yoki yaratish)
    current_dict = get_current_dict(uid)
    
    is_new = not any(w["ko"] == word for w in current_dict)
    
    if is_new:
        # Yangi so'zni faqat joriy tarjima tilidagi lug'atga yozish
        current_dict.append({"ko": word, "uz": tr}) 
        
        # O'zgarishlarni users obyektiga qayta yozish
        if 'dictionaries' not in users[uid]: users[uid]['dictionaries'] = {} # Himoya
        users[uid]['dictionaries'][tr_code] = current_dict
        save_users(users)
        record_history(uid, "add", f"{word} ({tr_code.upper()})") 

    text = f"🇰🇷 <b>{word}</b> -> {tr_code.upper()}\n"
    for i,x in enumerate(tr,1):
        text += f"{i}. {html.escape(x)}\n" 
    
    await m.answer(text, parse_mode="HTML", reply_markup=stop_kb(uid))


# ================= MYDICT (YANGILANDI) =================
@router.callback_query(F.data=="mydict")
async def mydict(cb:CallbackQuery, state:FSMContext):
    uid = str(cb.from_user.id)
    
    await state.set_state(MyDictState.active)

    current_dict = get_current_dict(uid)
    tr_code = tr_lang(uid)

    if not current_dict:
        await cb.message.edit_text(t(uid,"empty"), reply_markup=start_menu_kb_inline(uid))
        await cb.answer()
        return

    word_list_text = t(uid, "word_list", tr_lang_code=tr_code.upper())
    
    for idx, w in enumerate(current_dict, 1):
        translations = ", ".join([x.replace('*', '').replace('_', '') for x in w.get("uz", [])])
        word_list_text += f"**{idx}.** 🇰🇷 **{w['ko']}** -> {tr_code.upper()} {translations}\n"
    
    await cb.message.edit_text(
        word_list_text, 
        parse_mode="Markdown",
        reply_markup=mydict_management_kb(uid)
    )
    await cb.answer()


# ================= EDIT PROCESS (YANGILANDI) =================
@router.callback_query(F.data=="start_edit", MyDictState.active)
async def start_edit(cb:CallbackQuery, state:FSMContext):
    uid = str(cb.from_user.id)
    tr_code = tr_lang(uid).upper()
    await state.set_state(EditState.waiting_word)
    # tr_lang_code ni matnga qo'shdik
    await cb.message.edit_text(t(uid,"edit_prompt", tr_lang_code=tr_code), parse_mode="Markdown")
    await cb.answer()

@router.message(EditState.waiting_word)
async def process_edit(m:Message, state:FSMContext):
    uid = str(m.from_user.id)
    tr_code = tr_lang(uid)
    text = m.text.strip()
    
    if '.' not in text:
        await m.answer(f"❌ Noto‘g‘ri format. Iltimos, `Koreys_so'z.tarjima` formatida kiriting.\n\n{t(uid,'edit_prompt', tr_lang_code=tr_code.upper())}", parse_mode="Markdown")
        return

    try:
        ko, translations_str = text.split('.', 1)
        ko = ko.strip()
        new_tr = [x.strip() for x in translations_str.split(',') if x.strip()]
        if not ko or not new_tr: raise ValueError
    except ValueError:
        await m.answer(f"❌ Xatolik. Qaytadan urinib ko‘ring.\n\n{t(uid,'edit_prompt', tr_lang_code=tr_code.upper())}", parse_mode="Markdown")
        return

    users = load_users()
    # Faqat joriy tanlangan tildagi lug'atni tahrirlaymiz
    current_dict = users[uid]['dictionaries'].get(tr_code, []) 
    found = False
    
    for w in current_dict:
        if w["ko"] == ko:
            w["uz"] = new_tr
            found = True
            break
            
    if found:
        # Yangilangan lug'atni saqlash
        users[uid]['dictionaries'][tr_code] = current_dict 
        save_users(users)
        record_history(uid, "edit", f"{ko} ({tr_code.upper()})") 
        
        await m.answer(t(uid, "edit_success_back_to_mydict"))
        
        # Lug'atni qayta yuklash uchun Fake Callback
        class FakeCallback:
            def __init__(self, from_user, message):
                self.from_user = from_user
                self.message = message
                self.data = "mydict"
                
            async def answer(self, text=None, show_alert=False):
                pass 
                
        fake_cb = FakeCallback(m.from_user, m)
        await mydict(fake_cb, state) 
    else:
        await m.answer(t(uid,"not_found"), reply_markup=back_to_menu_kb(uid))
        
    await state.clear()


# ================= DELETE PROCESS (YANGILANDI) =================
@router.callback_query(F.data=="start_delete", MyDictState.active)
async def start_delete(cb:CallbackQuery, state:FSMContext):
    uid = str(cb.from_user.id)
    tr_code = tr_lang(uid).upper()
    await state.set_state(DeleteState.waiting_word)
    # tr_lang_code ni matnga qo'shdik
    await cb.message.edit_text(t(uid,"delete_prompt", tr_lang_code=tr_code), parse_mode="Markdown")
    await cb.answer()

@router.message(DeleteState.waiting_word)
async def process_delete(m:Message, state:FSMContext):
    uid = str(m.from_user.id)
    tr_code = tr_lang(uid)
    word = m.text.strip()
    
    users = load_users()
    # Faqat joriy tanlangan tildagi lug'atni o'chiramiz
    current_dict = users[uid]['dictionaries'].get(tr_code, [])
    initial_len = len(current_dict)
    
    new_dict = [w for w in current_dict if w["ko"] != word]
    
    if len(new_dict) < initial_len:
        # Yangilangan lug'atni saqlash
        users[uid]['dictionaries'][tr_code] = new_dict 
        save_users(users)
        record_history(uid, "delete", f"{word} ({tr_code.upper()})") 
        
        await m.answer(t(uid, "delete_success"))
        
        class FakeCallback:
            def __init__(self, from_user, message):
                self.from_user = from_user
                self.message = message
                self.data = "mydict"
                
            async def answer(self, text=None, show_alert=False):
                pass 
                
        fake_cb = FakeCallback(m.from_user, m)
        await mydict(fake_cb, state) 
        
    else:
        await m.answer(t(uid,"not_found"), reply_markup=back_to_menu_kb(uid))
        
    await state.clear()


# ================= SETTINGS (O'zgarishsiz qoldi) =================
@router.callback_query(F.data=="settings")
async def settings(cb:CallbackQuery, state: FSMContext):
    uid = str(cb.from_user.id)
    await state.clear()
    await cb.message.edit_text(t(uid,"settings_menu"), reply_markup=settings_kb(uid))
    await cb.answer()

@router.callback_query(F.data=="open_ui_lang_settings")
async def open_ui_lang_settings(cb:CallbackQuery):
    await cb.message.edit_text(t(cb.from_user.id, "lang_menu"), reply_markup=lang_kb(prefix="change_uilang"))
    await cb.answer()

@router.callback_query(F.data.startswith("change_uilang_"))
async def change_ui_lang(cb: CallbackQuery):
    uid = str(cb.from_user.id)
    users = load_users()
    
    new_lang = cb.data.split("_")[2]
    
    users[uid]["lang"] = new_lang
    save_users(users)
    
    await cb.message.edit_text(t(uid,"settings_menu"), reply_markup=settings_kb(uid))
    await cb.answer(text="✅ Boshqaruv tili o‘zgartirildi.") 

@router.callback_query(F.data=="open_tr_lang_settings")
async def open_tr_lang_settings(cb:CallbackQuery):
    uid = str(cb.from_user.id)
    await cb.message.edit_text(t(uid, "tr_lang_menu"), reply_markup=tr_lang_kb(uid)) 
    await cb.answer()

@router.callback_query(F.data.startswith("trlang_"))
async def change_tr_lang(cb: CallbackQuery):
    uid = str(cb.from_user.id)
    users = load_users()
    
    new_tr_lang = cb.data.split("_")[1]
    
    users[uid]["tr_lang"] = new_tr_lang
    save_users(users)
    
    await cb.message.edit_text(t(uid,"settings_menu"), reply_markup=settings_kb(uid))
    await cb.answer(text="✅ Tarjima tili o‘zgartirildi.") 

# ================= DOWNLOAD HANDLERS (YANGILANDI) =================
@router.callback_query(F.data.startswith("export_"), DownloadState.waiting_format)
async def process_export(cb: CallbackQuery, state: FSMContext):
    uid = str(cb.from_user.id)
    export_format = cb.data.split("_")[1]
    
    await state.clear()
    
    current_tr_lang = tr_lang(uid)
    # Faqat joriy tanlangan tildagi lug'atni olish
    words_to_export = get_current_dict(uid) 
        
    if not words_to_export:
        await cb.message.edit_text(t(uid,"empty"), reply_markup=start_menu_kb_inline(uid))
        await cb.answer()
        return

    file_path = None
    try:
        if export_format == "docx":
            file_path = create_docx(words_to_export, uid, current_tr_lang)
        elif export_format == "xlsx":
            file_path = create_xlsx(words_to_export, uid, current_tr_lang)
        elif export_format == "json":
            file_path = create_json(words_to_export, uid, current_tr_lang)
        
        if file_path:
            document = FSInputFile(file_path)
            await bot.send_document(cb.message.chat.id, document, caption=t(uid, "file_sent"))
            os.remove(file_path) 
            
    except Exception as e:
        await cb.message.answer(f"❌ Faylni yuklashda xatolik yuz berdi: {e}", reply_markup=start_menu_kb_inline(uid))
    
    await cb.message.delete()
    await cb.message.answer(t(uid, "menu"), reply_markup=start_menu_kb_inline(uid))
    await cb.answer()

# ================= ADMIN PANEL HANDLERS (STATISTIKA YANGILANDI) =================
@router.callback_query(F.data=="stats")
async def view_stats(cb: CallbackQuery, state: FSMContext):
    uid = str(cb.from_user.id)
    await state.clear()
    users = load_users()
    user_data = users.get(uid, {})
    
    # Yangilangan hisoblash:
    all_words_flat = get_all_words_flat(uid)
    total_words = len(all_words_flat) # Barcha tillardagi jami so'zlar
    
    current_dict = get_current_dict(uid)
    current_words = len(current_dict) # Joriy tildagi so'zlar
    current_tr_code = tr_lang(uid).upper()

    history = user_data.get("history", [])
    today = date.today().strftime("%Y-%m-%d")
    
    today_actions = {"add": 0, "edit": 0, "delete": 0}
    
    for item in history:
        item_date = item.get("time", "")[:10]
        if item_date == today:
            action = item.get("action")
            # Faqat joriy lug'atga tegishli harakatlarni hisoblash uchun word ichida tr_code.upper() ni tekshiramiz
            if current_tr_code in item.get("word", "").upper():
                if action in today_actions:
                    today_actions[action] += 1

    stats_text = t(uid, "stats_text", 
        total_words=total_words,
        current_words=current_words, 
        tr_lang_code=current_tr_code,
        today_date=today,
        added=today_actions["add"],
        edited=today_actions["edit"],
        deleted=today_actions["delete"]
    )
    
    await cb.message.edit_text(stats_text, parse_mode="Markdown", reply_markup=back_to_menu_kb(uid))
    await cb.answer()

# ... admin panelning qolgan qismi o'zgarishsiz qoldi ...

# ================= GAME HANDLERS (YANGILANDI) =================
@router.callback_query(F.data=="game_menu")
async def game_menu(cb: CallbackQuery, state: FSMContext):
    uid = str(cb.from_user.id)
    await state.clear()
    await cb.message.edit_text(t(uid, "game"), reply_markup=game_menu_kb(uid))
    await cb.answer()

@router.callback_query(F.data=="start_game_auto")
async def start_game_auto(cb: CallbackQuery, state: FSMContext):
    uid = str(cb.from_user.id)
    
    words_list = get_current_dict(uid) # Faqat joriy lug'at

    if not words_list or len(words_list) < WORDS_PER_GAME:
        await cb.answer(f"❌ O‘yinni boshlash uchun joriy ({tr_lang(uid).upper()}) lug‘atda kamida {WORDS_PER_GAME} ta so‘z bo‘lishi kerak.", show_alert=True)
        return
        
    await state.set_state(GameState.playing)
    await state.set_data({"last_game_time": time.time() - GAME_INTERVAL + 60, "game_task": None}) 
    
    tr_code = tr_lang(uid).upper()
    await cb.message.edit_text(t(uid, "game_start", tr_lang_code=tr_code, interval_min=GAME_INTERVAL//60, words_count=WORDS_PER_GAME), reply_markup=game_menu_kb(uid))
    
    game_task = asyncio.create_task(periodic_game_check(uid, cb.message.chat.id))
    current_state_data = await state.get_data()
    current_state_data["game_task"] = game_task 
    await state.set_data(current_state_data)

    await cb.answer()

# ... stop_game_auto (O'zgarishsiz qoldi) ...
@router.callback_query(F.data=="stop_game_auto")
async def stop_game_auto(cb: CallbackQuery, state: FSMContext):
    uid = str(cb.from_user.id)
    state_data = await state.get_data()
    game_task = state_data.get("game_task")
    
    if game_task and not game_task.done():
        game_task.cancel()
        
    await state.clear()
    await cb.message.edit_text(t(uid, "game_stopped"), reply_markup=game_menu_kb(uid))
    await cb.answer()

async def periodic_game_check(uid, chat_id):
    """Har 12 daqiqada o'yin so'zlarini yuborish uchun doimiy tekshiruv."""
    while True:
        try:
            await asyncio.sleep(60) 
            
            state = FSMContext(storage=dp.storage, key=dp.storage.build_key(chat_id, uid))
            current_state = await state.get_state()
            state_data = await state.get_data()
            
            if current_state != GameState.playing:
                return 
                
            last_game_time = state_data.get("last_game_time", 0)
            
            if time.time() - last_game_time >= GAME_INTERVAL:
                await start_new_game(uid, chat_id, state)
                state_data["last_game_time"] = time.time()
                await state.set_data(state_data)
                
        except asyncio.CancelledError:
            return
        except Exception as e:
            print(f"Periodic game check xatosi: {e}")
            await asyncio.sleep(60)

async def start_new_game(uid, chat_id, state):
    """Foydalanuvchiga yangi o'yin so'zlarini yuboradi."""
    
    words = get_current_dict(uid) # Faqat joriy lug'atdagi so'zlar
    tr_code = tr_lang(uid)
    
    if len(words) < WORDS_PER_GAME:
        # Agar so'zlar soni kamayib ketsa, o'yinni to'xtatish haqida xabar berish mumkin.
        return 
        
    selected_words = random.sample(words, WORDS_PER_GAME)
    
    prompt_text = t(uid, "game_prompt", tr_lang_code=tr_code.upper()) + "\n\n"
    game_data = {}
    
    for idx, w in enumerate(selected_words, 1):
        translations = ", ".join([x.replace('*', '').replace('_', '') for x in w.get("uz", [])])
        prompt_text += f"{idx}. 🇰🇷 **{w['ko']}**\n"
        game_data[w['ko']] = f"{tr_code.upper()} {translations}" 
        
    await state.update_data({"current_game_words": game_data})
    
    try:
        await bot.send_message(
            chat_id, 
            prompt_text, 
            parse_mode="Markdown", 
            reply_markup=game_reveal_kb(uid)
        )
    except Exception as e:
        print(f"Xabar yuborishda xato: {e}")


@router.callback_query(F.data=="reveal_game_auto", GameState.playing)
async def reveal_game_auto(cb: CallbackQuery, state: FSMContext):
    uid = str(cb.from_user.id)
    state_data = await state.get_data()
    game_data = state_data.get("current_game_words", {})
    
    if not game_data:
        await cb.answer("❌ Tarjima yo‘q. Yangi o‘yinni kuting.", show_alert=True)
        return
        
    reveal_text = t(uid, "game_reveal") + "\n\n"
    
    for ko_word, tr_text in game_data.items():
        reveal_text += f"🇰🇷 **{ko_word}** -> {tr_text}\n"
        
    await cb.message.edit_text(reveal_text, parse_mode="Markdown", reply_markup=game_menu_kb(uid))
    await state.update_data({"current_game_words": {}}) 
    await cb.answer()

# ================= POLLING =================
async def main():
    try:
        await dp.start_polling(bot)
    finally:
        await bot.session.close()

if __name__ == "__main__":
    asyncio.run(main())