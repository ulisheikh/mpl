# -*- coding: utf-8 -*-
"""
KOREAN-O'ZBEK LUG'AT BOT - 1-QISM
IMPORT VA ASOSIY FUNKSIYALAR
"""

# ============================================
# IMPORT QISMI
# ============================================
import telebot
import json
import os
import re
import subprocess
import psutil
import threading
import time
import random
from datetime import datetime
from html import escape
from telebot import types
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill

# ============================================
# KONFIGURATSIYA
# ============================================
TOKEN = os.getenv("BOT_TOKEN", "8046756811:AAEsMXNBMkIMkqM3XtVyQ3OzOd4itRfn03M")
DATA_FILE = "dictionary.json"
START_TIME = datetime.now()
ADMIN_ID = 8046330769

bot = telebot.TeleBot(TOKEN)
user_context = {}
exam_mode = {}

# ============================================
# MA'LUMOTLAR BILAN ISHLASH
# ============================================

def load_data():
    if os.path.exists(DATA_FILE):
        try:
            with open(DATA_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            updated = False
            new_data = {}
            for key, value in data.items():
                if key.startswith("Topic-"):
                    new_key = key.replace("Topic-", "Topik-")
                    new_data[new_key] = value
                    updated = True
                else:
                    new_data[key] = value
            if updated:
                save_data(new_data)
                return new_data
            return data
        except: 
            return {}
    return {}

def save_data(data):
    try:
        if os.path.exists(DATA_FILE):
            backup_file = DATA_FILE + ".backup"
            with open(DATA_FILE, 'r', encoding='utf-8') as f:
                backup_data = f.read()
            with open(backup_file, 'w', encoding='utf-8') as f:
                f.write(backup_data)
        with open(DATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
    except Exception as e:
        print(f"Save error: {e}")

def is_korean(text):
    return bool(re.search('[\uac00-\ud7af]', text))

# ============================================
# IMTIHON FAYLI FUNKSIYALARI
# ============================================

def get_20_words_from_db(shuffle=False):
    data = load_data()
    all_words = []
    for topic in data.values():
        for section in topic.values():
            for question in section.values():
                for korean, uzbek in question.items():
                    all_words.append({'korean': korean, 'uzbek': uzbek})
    if len(all_words) < 20:
        return all_words
    if shuffle:
        selected = random.sample(all_words, 20)
    else:
        selected = all_words[:20]
    return selected

def create_docx_exam(words, filename="imtihon.docx"):
    doc = Document()
    title = doc.add_heading('🇰🇷 Koreys tili Imtihon varaqasi', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = '№'
    header_cells[1].text = "O'zbekcha so'z"
    header_cells[2].text = "Koreyscha so'z"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, word in enumerate(words, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = word['uzbek']
        row_cells[2].text = ''
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.save(filename)
    return filename

def create_excel_exam(words, filename="imtihon.xlsx"):
    wb = Workbook()
    ws = wb.active
    ws.title = "Imtihon"
    ws.merge_cells('A1:C1')
    ws['A1'] = '🇰🇷 Koreys tili Imtihon varaqasi'
    ws['A1'].font = Font(size=16, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    headers = ['№', "O'zbekcha so'z", "Koreyscha so'z"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=header)
        cell.font = Font(bold=True, size=12)
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
    for idx, word in enumerate(words, 1):
        ws.cell(row=idx+2, column=1, value=idx).alignment = Alignment(horizontal='center')
        ws.cell(row=idx+2, column=2, value=word['uzbek'])
        ws.cell(row=idx+2, column=3, value='')
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 30
    wb.save(filename)
    return filename

def get_main_keyboard():
    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.row("/start")
    markup.row("📂 BO'LIMLAR")
    markup.row("📥 JSON", "🐍 PYTHON")
    markup.row("📝 IMTIHON FAYLI")
    return markup

# ============================================
# KOMANDA HANDLERLAR
# ============================================

@bot.message_handler(commands=['start'])
def welcome_cmd(message):
    help_text = (
        "<b>📚 LUG'AT BOT | ADMIN</b>\n"
        "<b>───────────────────</b>\n\n"
        "<b>1️⃣ BO'LIM YARATISH</b>\n👉 <code>>35 r,w,l</code>\n\n"
        "<b>2️⃣ SO'Z QO'SHISH</b>\n👉 <code>1</code>\n👉 <code>안녕 salom</code>\n\n"
        "<b>3️⃣ O'CHIRISH</b>\n🗑 <code>rm.35r33</code>\n\n"
        "<b>4️⃣ QIDIRUV</b>\n🔍 <code>s.so'z</code>\n\n"
        "<b>⚙️ TIZIM:</b> /status"
    )
    bot.send_message(message.chat.id, help_text, reply_markup=get_main_keyboard(), parse_mode="HTML")

@bot.message_handler(func=lambda m: m.text == "📝 IMTIHON FAYLI")
def exam_file_menu(message):
    uid = message.from_user.id
    if uid not in exam_mode:
        exam_mode[uid] = 'tartibli'
    markup = types.InlineKeyboardMarkup(row_width=2)
    mode_text = "🔄 Mode: Tartibli" if exam_mode[uid] == 'tartibli' else "🔄 Mode: Aralash"
    mode_btn = types.InlineKeyboardButton(mode_text, callback_data='toggle_mode')
    docx_btn = types.InlineKeyboardButton("📄 DOCX", callback_data='exam_docx')
    excel_btn = types.InlineKeyboardButton("📊 EXCEL", callback_data='exam_excel')
    markup.add(mode_btn)
    markup.row(docx_btn, excel_btn)
    bot.send_message(
        message.chat.id,
        "📥 <b>Imtihon faylini yuklab olish:</b>\n\n"
        "1️⃣ Rejimni tanlang\n"
        "2️⃣ Format tanlang\n\n"
        "ℹ️ Har bir faylda 20 ta so'z.",
        reply_markup=markup,
        parse_mode="HTML"
    )

@bot.callback_query_handler(func=lambda call: call.data == 'toggle_mode')
def toggle_exam_mode(call):
    uid = call.from_user.id
    exam_mode[uid] = 'aralash' if exam_mode.get(uid) == 'tartibli' else 'tartibli'
    markup = types.InlineKeyboardMarkup(row_width=2)
    mode_text = "🔄 Mode: Tartibli" if exam_mode[uid] == 'tartibli' else "🔄 Mode: Aralash"
    mode_btn = types.InlineKeyboardButton(mode_text, callback_data='toggle_mode')
    docx_btn = types.InlineKeyboardButton("📄 DOCX", callback_data='exam_docx')
    excel_btn = types.InlineKeyboardButton("📊 EXCEL", callback_data='exam_excel')
    markup.add(mode_btn)
    markup.row(docx_btn, excel_btn)
    bot.edit_message_reply_markup(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        reply_markup=markup
    )
    bot.answer_callback_query(call.id, f"✅ Rejim: {exam_mode[uid].capitalize()}")

@bot.callback_query_handler(func=lambda call: call.data in ['exam_docx', 'exam_excel'])
def generate_exam_file(call):
    uid = call.from_user.id
    file_type = call.data.replace('exam_', '')
    shuffle = exam_mode.get(uid, 'tartibli') == 'aralash'
    words = get_20_words_from_db(shuffle=shuffle)
    if not words:
        bot.answer_callback_query(call.id, "❌ Bazada so'zlar yo'q!", show_alert=True)
        return
    bot.answer_callback_query(call.id, "⏳ Fayl tayyorlanmoqda...")
    try:
        if file_type == 'docx':
            filename = create_docx_exam(words)
            caption = "📄 Imtihon fayli (DOCX)"
        else:
            filename = create_excel_exam(words)
            caption = "📊 Imtihon fayli (EXCEL)"
        with open(filename, 'rb') as f:
            bot.send_document(call.message.chat.id, f, caption=caption)
        os.remove(filename)
    except Exception as e:
        bot.send_message(call.message.chat.id, f"❌ Xatolik: {e}")

# DAVOMINI 2-QISMDA KIRITING (barcha qolgan handlerlar)
# -*- coding: utf-8 -*-
"""
KOREAN-O'ZBEK LUG'AT BOT - 2-QISM
BARCHA HANDLERLAR (1-qismdan keyin qo'shiladi)
"""

# ============================================
# ASOSIY MATN HANDLER
# ============================================

@bot.message_handler(content_types=['text'])
def handle_all(message):
    text = message.text.strip().lower()
    
    if text.startswith("/"):
        return
    
    uid = message.from_user.id
    data = load_data()
    
    if uid not in user_context:
        user_context[uid] = {}

    # HOZIRGI JOY (%l)
    if text == "%l":
        topic = user_context[uid].get("topic", "Tanlanmagan")
        section = user_context[uid].get("section", "Tanlanmagan")
        question = user_context[uid].get("question", "Tanlanmagan")
        bot.send_message(
            message.chat.id, 
            f"📍 Hozirgi manzilingiz:\n\n"
            f"📂 Topik: {topic}\n"
            f"📖 Bo'lim: {section}\n"
            f"🔢 Savol: {question}"
        )
        return

    # TIKLASH (rs.)
    if text.lower().startswith("rs."):
        target = text[3:].strip()
        
        # Topikni tiklash (rs.35)
        if target.isdigit():
            topic_num = int(target)
            backup_file = f"backup_topik_{topic_num}.json"
            if not os.path.exists(backup_file):
                return bot.send_message(message.chat.id, f"❌ {topic_num}-topik uchun backup topilmadi")
            try:
                with open(backup_file, 'r', encoding='utf-8') as f:
                    topic_data = json.load(f)
                topic_key = f"Topik-{topic_num}"
                data[topic_key] = topic_data
                save_data(data)
                os.remove(backup_file)
                bot.send_message(message.chat.id, f"✅ {topic_num}-topik tiklandi!")
            except Exception as e:
                bot.send_message(message.chat.id, f"❌ Xatolik: {e}")
            return
        
        # Bo'limni tiklash (rs.35r)
        if len(target) >= 2 and target[:-1].isdigit() and target[-1] in ['r', 'w', 'l']:
            topic_num = int(target[:-1])
            section_code = target[-1]
            section_map = {'r': 'reading', 'w': 'writing', 'l': 'listening'}
            section_name = section_map[section_code]
            backup_file = f"backup_{topic_num}_{section_name}.json"
            if not os.path.exists(backup_file):
                return bot.send_message(message.chat.id, f"❌ Backup topilmadi")
            try:
                with open(backup_file, 'r', encoding='utf-8') as f:
                    section_data = json.load(f)
                topic_key = f"Topik-{topic_num}"
                if topic_key not in data:
                    data[topic_key] = {}
                data[topic_key][section_name] = section_data
                save_data(data)
                os.remove(backup_file)
                bot.send_message(message.chat.id, f"✅ Tiklandi!")
            except Exception as e:
                bot.send_message(message.chat.id, f"❌ Xatolik: {e}")
            return

        # Savolni tiklash (rs.35r33)
        if any(c in target for c in ['r', 'w', 'l']) and target[-1].isdigit() and len(target) > 3:
            backup_file = f"backup_q_{target}.json"
            if os.path.exists(backup_file):
                try:
                    with open(backup_file, 'r', encoding='utf-8') as f:
                        b = json.load(f)
                    t_k, s_k, q_k = b['topic'], b['section'], b['question_key']
                    if t_k not in data: data[t_k] = {}
                    if s_k not in data[t_k]: data[t_k][s_k] = {}
                    data[t_k][s_k][q_k] = b['content']
                    save_data(data)
                    os.remove(backup_file)
                    bot.send_message(message.chat.id, f"✅ Savol tiklandi!")
                except Exception as e:
                    bot.send_message(message.chat.id, f"❌ Xato: {e}")
                return

        # So'zni tiklash
        safe_name = "".join(x for x in target if x.isalnum())
        backup_file = f"backup_word_{safe_name}.json"
        if os.path.exists(backup_file):
            try:
                with open(backup_file, 'r', encoding='utf-8') as f:
                    b = json.load(f)
                t_k, s_k, q_k = b['topic'], b['section'], b['question']
                kor, uz = b['korean'], b['uzbek']
                if t_k not in data: data[t_k] = {}
                if s_k not in data[t_k]: data[t_k][s_k] = {}
                if q_k not in data[t_k][s_k]: data[t_k][s_k][q_k] = {}
                data[t_k][s_k][q_k][kor] = uz
                save_data(data)
                os.remove(backup_file)
                bot.send_message(message.chat.id, f"✅ So'z tiklandi!\n\n📍 {t_k} → {s_k} → {q_k}\n🇰🇷 {kor} = 🇺🇿 {uz}")
            except Exception as e:
                bot.send_message(message.chat.id, f"❌ Xatolik: {e}")
            return

    # O'CHIRISH (rm.)
    if text.lower().startswith("rm."):
        target = text[3:].strip()
        
        # Bo'limni o'chirish
        if len(target) >= 2 and target[:-1].isdigit() and target[-1] in ['r', 'w', 'l']:
            topic_num = int(target[:-1])
            section_code = target[-1]
            topic_key = f"Topik-{topic_num}"
            section_map = {'r': 'reading', 'w': 'writing', 'l': 'listening'}
            section_name = section_map[section_code]
            if topic_key in data and section_name in data[topic_key]:
                backup_file = f"backup_{topic_num}_{section_name}.json"
                with open(backup_file, 'w', encoding='utf-8') as f:
                    json.dump(data[topic_key][section_name], f, ensure_ascii=False, indent=4)
                del data[topic_key][section_name]
                save_data(data)
                bot.send_message(message.chat.id, f"🗑 O'chirildi\n\n💡 Qaytarish: rs.{topic_num}{section_code}")
            else:
                bot.send_message(message.chat.id, "❌ Topilmadi")
            return
        
        # Topikni o'chirish
        if target.isdigit():
            topic_num = int(target)
            topic_key = f"Topik-{topic_num}"
            if topic_key in data:
                backup_file = f"backup_topik_{topic_num}.json"
                with open(backup_file, 'w', encoding='utf-8') as f:
                    json.dump(data[topic_key], f, ensure_ascii=False, indent=4)
                del data[topic_key]
                save_data(data)
                bot.send_message(message.chat.id, f"🗑 {topic_num}-topik o'chirildi\n\n💡 Qaytarish: rs.{topic_num}")
            else:
                bot.send_message(message.chat.id, f"❌ Topilmadi")
            return

        # Savolni o'chirish
        pattern_q = re.match(r'^(\d+)([rwl])(\d+)$', target)
        if pattern_q:
            topic_num = int(pattern_q.group(1))
            section_code = pattern_q.group(2)
            q_num = pattern_q.group(3)
            topic_key = f"Topik-{topic_num}"
            section_map = {'r': 'reading', 'w': 'writing', 'l': 'listening'}
            section_name = section_map[section_code]
            q_key = f"{q_num}-savol so'zlari"
            if topic_key in data and section_name in data[topic_key] and q_key in data[topic_key][section_name]:
                backup_info = {
                    'type': 'question',
                    'topic': topic_key,
                    'section': section_name,
                    'question_key': q_key,
                    'content': data[topic_key][section_name][q_key]
                }
                backup_file = f"backup_q_{target}.json"
                with open(backup_file, 'w', encoding='utf-8') as f:
                    json.dump(backup_info, f, ensure_ascii=False, indent=4)
                del data[topic_key][section_name][q_key]
                save_data(data)
                bot.send_message(message.chat.id, f"🗑 Savol o'chirildi\n\n💡 Qaytarish: rs.{target}")
            else:
                bot.send_message(message.chat.id, "❌ Topilmadi")
            return

        # So'zni o'chirish
        word_to_rm = target.lower()
        found = False
        for t in data:
            for s in data[t]:
                for q in data[t][s]:
                    match_key = None
                    if word_to_rm in [k.lower() for k in data[t][s][q].keys()]:
                        match_key = next(k for k in data[t][s][q].keys() if k.lower() == word_to_rm)
                    elif word_to_rm in [v.lower() for v in data[t][s][q].values()]:
                        match_key = next(k for k, v in data[t][s][q].items() if v.lower() == word_to_rm)
                    if match_key:
                        backup_info = {
                            'type': 'word', 'topic': t, 'section': s, 'question': q,
                            'korean': match_key, 'uzbek': data[t][s][q][match_key]
                        }
                        safe_name = "".join(x for x in word_to_rm if x.isalnum())
                        backup_file = f"backup_word_{safe_name}.json"
                        with open(backup_file, 'w', encoding='utf-8') as f:
                            json.dump(backup_info, f, ensure_ascii=False, indent=4)
                        del data[t][s][q][match_key]
                        save_data(data)
                        msg = f"🗑 So'z o'chirildi: <b>{match_key}</b>\n\n📍 {t} → {s} → {q}\n💡 Qaytarish: <code>rs.{word_to_rm}</code>"
                        bot.send_message(message.chat.id, msg, parse_mode="HTML")
                        found = True
                        break
                if found: break
            if found: break
        if not found:
            bot.send_message(message.chat.id, "❌ Topilmadi")
        return

    # BO'LIM YARATISH (>35r1)
    if text.startswith(">"):
        match = re.match(r">(\d+)([rwl,]+)?(\d+)?(\.)?", text)
        if not match:
            return bot.send_message(message.chat.id, "❌ Noto'g'ri format!")
        t_num, sec_codes, q_num, view_mode = match.groups()
        topic_key = f"Topik-{t_num}"
        section_map = {'r': 'reading', 'w': 'writing', 'l': 'listening'}
        
        # Ko'rish rejimi (>35.)
        if view_mode == ".":
            if topic_key not in data:
                return bot.send_message(message.chat.id, f"❌ {topic_key} yo'q")
            if sec_codes:
                sec_name = section_map.get(sec_codes[0])
                section_data = data.get(topic_key, {}).get(sec_name, {})
                if not section_data:
                    return bot.send_message(message.chat.id, f"⚠️ {sec_name.upper()} bo'sh")
                try:
                    def extract_num(s):
                        found = re.findall(r'\d+', s)
                        return int(found[0]) if found else 0
                    sorted_keys = sorted(section_data.keys(), key=extract_num)
                except:
                    sorted_keys = sorted(section_data.keys())
                full_message = f"📌 <b>{topic_key} > {sec_name.upper()}</b>\n\n"
                for q_key in sorted_keys:
                    words_dict = section_data[q_key]
                    block = f"<code>{q_key}</code>\n"
                    for kr, uz in words_dict.items():
                        block += f"  • <b>{kr}</b> - <i>{uz}</i>\n"
                    block += "\n"
                    if len(full_message + block) > 4000:
                        bot.send_message(message.chat.id, full_message, parse_mode="HTML")
                        full_message = block
                    else:
                        full_message += block
                if full_message:
                    bot.send_message(message.chat.id, full_message, parse_mode="HTML")
            else:
                secs = list(data.get(topic_key, {}).keys())
                bot.send_message(message.chat.id, f"📚 <b>{topic_key}</b> bo'limlari:\n\n" + "\n".join([f"• {s.upper()}" for s in secs]), parse_mode="HTML")
            return

        # Yaratish rejimi
        is_new = False
        if topic_key not in data:
            data[topic_key] = {}
            is_new = True
        user_context[uid]["topic"] = topic_key
        
        active_sec = None
        if sec_codes:
            codes = sec_codes.replace(",", "")
            for c in codes:
                s_name = section_map.get(c)
                if s_name:
                    if s_name not in data[topic_key]:
                        data[topic_key][s_name] = {}
                        is_new = True
                    if not active_sec: active_sec = s_name
            user_context[uid]["section"] = active_sec
        
        if q_num:
            q_key = f"{q_num}-savol so'zlari"
            if active_sec:
                if q_key not in data[topic_key][active_sec]:
                    data[topic_key][active_sec][q_key] = {}
                    is_new = True
                user_context[uid]["question"] = q_key
        
        if active_sec and topic_key in data and active_sec in data[topic_key]:
            sorted_questions = dict(sorted(
                data[topic_key][active_sec].items(),
                key=lambda x: int(re.search(r'\d+', x[0]).group()) if re.search(r'\d+', x[0]) else 0
            ))
            data[topic_key][active_sec] = sorted_questions
        
        save_data(data)
        status = "yaratildi ✨" if is_new else "tanlandi ✅"
        loc = f"📍 {topic_key}"
        if active_sec: loc += f" > {active_sec}"
        if q_num: loc += f" > {q_num}-savol"
        bot.send_message(message.chat.id, f"{loc} {status}")
        return

    # QIDIRISH (s.so'z)
    if text.startswith("s."):
        query = text[2:].strip().lower()
        results = []
        for t, s_dict in data.items():
            topic_num = t.replace("Topik-", "")
            for s, q_dict in s_dict.items():
                for q, w_dict in q_dict.items():
                    q_num = q.replace("-savol so'zlari", "")
                    for kr, uz in w_dict.items():
                        if query in kr.lower() or query in uz.lower():
                            results.append(f"📍 {topic_num}-topik > {s.upper()} > {q_num}-savol: {kr} → {uz}")
        if results:
            msg = "🔍 TOPILDI:\n\n" + "\n\n".join(results)
        else:
            msg = "❌ Topilmadi"
        bot.send_message(message.chat.id, msg)
        return

    # SO'Z QO'SHISH
    if " " in text and not text.startswith(">") and not text.isdigit() and not text.startswith("s."):
        t = user_context[uid].get("topic")
        s = user_context[uid].get("section")
        q = user_context[uid].get("question")
        if not t or not s or not q:
            return bot.send_message(message.chat.id, "⚠️ Avval manzilni kiriting! Misol: >35r1")
        lines = message.text.strip().split('\n')
        added_count = 0
        added_words_list = []
        for line in lines:
            line = line.strip()
            if not line: continue
            split_match = re.search(r'([가-힣])\s+([a-zA-Z])|([a-zA-Z])\s+([가-힣])', line)
            if split_match:
                idx = split_match.start(0) + 1
                part1 = line[:idx].strip()
                part2 = line[idx:].strip()
                if re.search(r'[가-힣]', part1):
                    kor, uzb = part1, part2
                else:
                    kor, uzb = part2, part1
            else:
                parts = line.split(" ", 1)
                if len(parts) == 2:
                    p1, p2 = parts[0].strip(), parts[1].strip()
                    if re.search(r'[가-힣]', p1): kor, uzb = p1, p2
                    else: kor, uzb = p2, p1
                else:
                    continue
            if t not in data: data[t] = {}
            if s not in data[t]: data[t][s] = {}
            if q not in data[t][s]: data[t][s][q] = {}
            data[t][s][q][kor] = uzb
            added_words_list.append(f"• {kor} - {uzb}")
            added_count += 1
        if added_count > 0:
            save_data(data)
            clean_q = q.replace(" so'zlari", "")
            report = f"✅ **{added_count} ta so'z saqlandi!**\n\n"
            report += "\n".join(added_words_list)
            report += f"\n\n📍 **Manzil:** {t} > {s.upper()} > {clean_q}"
            bot.send_message(message.chat.id, report, parse_mode="Markdown")
        else:
            bot.send_message(message.chat.id, "❌ Format xato!")
        return

# ============================================
# ISHGA TUSHIRISH
# ============================================

if __name__ == "__main__":
    try:
        me = bot.get_me()
        print(f"\n{'='*40}")
        print(f"BOT: @{me.username}")
        print(f"STATUS: RUNNING ✅")
        print(f"{'='*40}\n")
        bot.infinity_polling()
    except Exception as e: 
        print(f"ERROR: {e}")